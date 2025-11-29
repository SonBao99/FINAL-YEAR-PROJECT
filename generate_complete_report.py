#!/usr/bin/env python3
"""
Generate complete report with ALL screenshots embedded
Finds all screenshots in the directory and matches them with report content
"""

import re
import base64
from pathlib import Path
from typing import Optional, Dict
from collections import defaultdict

def encode_image_to_base64(image_path: Path) -> Optional[str]:
    """Encode image file to base64 string"""
    if not image_path.exists():
        return None
    try:
        with open(image_path, 'rb') as f:
            image_data = f.read()
            base64_str = base64.b64encode(image_data).decode('utf-8')
            ext = image_path.suffix.lower()
            mime_types = {
                '.png': 'image/png',
                '.jpg': 'image/jpeg',
                '.jpeg': 'image/jpeg',
                '.gif': 'image/gif'
            }
            mime_type = mime_types.get(ext, 'image/png')
            return f"data:{mime_type};base64,{base64_str}"
    except Exception as e:
        print(f"Error encoding {image_path}: {e}")
        return None

def find_all_screenshots(screenshots_dir: Path) -> Dict[str, Path]:
    """Find all screenshot files and create a mapping"""
    screenshots = {}
    
    # Find all PNG files in code directory
    code_dir = screenshots_dir / 'code'
    if code_dir.exists():
        for png_file in code_dir.glob('fig-*.png'):
            # Extract figure number from filename
            # e.g., fig-6.35-Health-Check-Endpoint.png -> 6.35
            match = re.search(r'fig-(\d+\.\d+)-', png_file.name)
            if match:
                fig_num = match.group(1)
                screenshots[fig_num] = png_file
                screenshots[png_file.name] = png_file  # Also index by full name
    
    # Also check for UI screenshots
    ui_dir = screenshots_dir / 'ui'
    if ui_dir.exists():
        for png_file in ui_dir.glob('fig-*.png'):
            match = re.search(r'fig-(\d+\.\d+)-', png_file.name)
            if match:
                fig_num = match.group(1)
                screenshots[fig_num] = png_file
                screenshots[png_file.name] = png_file
    
    return screenshots

def generate_word_document_complete(report_path: Path, output_docx: Path, screenshots_dir: Path):
    """Generate Word document with ALL available screenshots"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("python-docx not installed. Installing...")
        import subprocess
        import sys
        subprocess.check_call([sys.executable, '-m', 'pip', 'install', 'python-docx'])
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    # Find all screenshots
    all_screenshots = find_all_screenshots(screenshots_dir)
    print(f"Found {len([k for k in all_screenshots.keys() if '.' in k])} screenshots by figure number")
    
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    with open(report_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    screenshot_count = 0
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # Check for screenshot placeholder
        screenshot_match = re.search(r'\[INSERT SCREENSHOT: ([^\]]+)\]', line)
        if screenshot_match:
            screenshot_path = screenshot_match.group(1)
            
            # Try multiple ways to find the screenshot
            screenshot_file = None
            
            # Method 1: Try exact path
            full_path = screenshots_dir.parent / screenshot_path
            if full_path.exists():
                screenshot_file = full_path
            
            # Method 2: Try just the filename
            if not screenshot_file:
                filename = Path(screenshot_path).name
                if filename in all_screenshots:
                    screenshot_file = all_screenshots[filename]
            
            # Method 3: Extract figure number and find by number
            if not screenshot_file:
                match = re.search(r'fig-(\d+\.\d+)-', screenshot_path)
                if match:
                    fig_num = match.group(1)
                    if fig_num in all_screenshots:
                        screenshot_file = all_screenshots[fig_num]
            
            # Method 4: Try in code directory
            if not screenshot_file:
                code_path = screenshots_dir / Path(screenshot_path).name
                if code_path.exists():
                    screenshot_file = code_path
            
            if screenshot_file and screenshot_file.exists():
                try:
                    paragraph = doc.add_paragraph()
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph.add_run()
                    run.add_picture(str(screenshot_file), width=Inches(6))
                    screenshot_count += 1
                    print(f"  [{screenshot_count}] Added: {screenshot_file.name}")
                except Exception as e:
                    p = doc.add_paragraph(f"[IMAGE ERROR: {screenshot_path} - {e}]")
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    print(f"  Could not add image {screenshot_file}: {e}")
            else:
                # Add placeholder
                p = doc.add_paragraph(f"[IMAGE PLACEHOLDER: {screenshot_path}]")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                print(f"  Image not found: {screenshot_path}")
            
            # Add caption (next line should be the caption)
            i += 1
            if i < len(lines) and lines[i].strip().startswith('Figure'):
                caption = lines[i].strip()
                p = doc.add_paragraph(caption)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.style = 'Caption'
                i += 1
            
            # Add explanation (next line should be explanation)
            if i < len(lines) and lines[i].strip() and not lines[i].strip().startswith('Figure'):
                explanation = lines[i].strip()
                if explanation:
                    p = doc.add_paragraph(explanation)
                    p.style = 'Normal'
                i += 1
            continue
        
        # Regular text
        if line:
            # Check if it's a heading
            if line.startswith('Section') or (line.startswith('Figure') and ':' in line):
                p = doc.add_heading(line, level=1 if line.startswith('Section') else 2)
            elif re.match(r'^\d+\.\d+', line):  # Subsection like "6.1"
                p = doc.add_heading(line, level=2)
            elif re.match(r'^\d+\.\d+\.\d+', line):  # Subsubsection like "6.1.1"
                p = doc.add_heading(line, level=3)
            else:
                # Regular paragraph
                p = doc.add_paragraph(line)
        
        i += 1
    
    doc.save(str(output_docx))
    print(f"\n✅ Word document generated: {output_docx}")
    print(f"✅ Total screenshots embedded: {screenshot_count}")
    print(f"✅ You can now open this file directly in Microsoft Word!")

def main():
    import argparse
    
    parser = argparse.ArgumentParser(
        description='Generate complete report with ALL embedded images'
    )
    parser.add_argument(
        '--report',
        type=str,
        default='report.txt',
        help='Input report file (default: report.txt)'
    )
    parser.add_argument(
        '--screenshots-dir',
        type=str,
        default='screenshots/code',
        help='Screenshots directory (default: screenshots/code)'
    )
    parser.add_argument(
        '--output',
        type=str,
        default='report_complete_with_images.docx',
        help='Output filename (default: report_complete_with_images.docx)'
    )
    
    args = parser.parse_args()
    
    report_path = Path(args.report)
    screenshots_dir = Path(args.screenshots_dir).parent  # Go up one level to screenshots/
    output_docx = Path(args.output)
    
    if not report_path.exists():
        print(f"Error: Report file not found: {report_path}")
        return
    
    print(f"Generating complete report...")
    print(f"Report file: {report_path}")
    print(f"Screenshots directory: {screenshots_dir}")
    print(f"Output file: {output_docx}")
    print("-" * 60)
    
    generate_word_document_complete(report_path, output_docx, screenshots_dir)

if __name__ == '__main__':
    main()

