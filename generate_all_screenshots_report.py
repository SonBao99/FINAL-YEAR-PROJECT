#!/usr/bin/env python3
"""
Generate report with ALL available screenshots embedded
Adds all screenshots from the directory, even if not referenced in report.txt
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def find_all_screenshots(screenshots_dir: Path):
    """Find all screenshot files and organize by figure number"""
    screenshots = {}
    code_dir = screenshots_dir / 'code'
    
    if code_dir.exists():
        for png_file in sorted(code_dir.glob('fig-*.png')):
            # Extract figure number: fig-6.35-Health-Check-Endpoint.png -> 6.35
            match = re.search(r'fig-(\d+\.\d+)-(.+)\.png', png_file.name)
            if match:
                fig_num = match.group(1)
                title = match.group(2).replace('-', ' ')
                screenshots[fig_num] = {
                    'path': png_file,
                    'filename': png_file.name,
                    'title': title,
                    'section': fig_num.split('.')[0]
                }
    
    return screenshots

def generate_complete_report(report_path: Path, output_docx: Path, screenshots_dir: Path):
    """Generate Word document with ALL screenshots"""
    
    # Find all screenshots
    all_screenshots = find_all_screenshots(screenshots_dir)
    print(f"Found {len(all_screenshots)} code screenshots")
    
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Read report content
    with open(report_path, 'r', encoding='utf-8') as f:
        report_lines = f.readlines()
    
    # Track which screenshots we've added
    added_screenshots = set()
    screenshot_count = 0
    
    # Process report and add screenshots where referenced
    i = 0
    while i < len(report_lines):
        line = report_lines[i].strip()
        
        # Check for screenshot placeholder
        screenshot_match = re.search(r'\[INSERT SCREENSHOT: ([^\]]+)\]', line)
        if screenshot_match:
            screenshot_path = screenshot_match.group(1)
            
            # Try to find the screenshot
            screenshot_file = None
            
            # Extract figure number from path
            match = re.search(r'fig-(\d+\.\d+)-', screenshot_path)
            if match:
                fig_num = match.group(1)
                if fig_num in all_screenshots:
                    screenshot_file = all_screenshots[fig_num]['path']
                    added_screenshots.add(fig_num)
            
            # Try alternative: just filename
            if not screenshot_file:
                filename = Path(screenshot_path).name
                for fig_num, info in all_screenshots.items():
                    if info['filename'] == filename:
                        screenshot_file = info['path']
                        added_screenshots.add(fig_num)
                        break
            
            # Try alternative: direct path
            if not screenshot_file:
                full_path = screenshots_dir.parent / screenshot_path
                if full_path.exists():
                    screenshot_file = full_path
            
            if screenshot_file and screenshot_file.exists():
                try:
                    paragraph = doc.add_paragraph()
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph.add_run()
                    run.add_picture(str(screenshot_file), width=Inches(6))
                    screenshot_count += 1
                    print(f"  [{screenshot_count}] Added: {screenshot_file.name}")
                except Exception as e:
                    p = doc.add_paragraph(f"[IMAGE ERROR: {e}]")
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    print(f"  Error adding {screenshot_file.name}: {e}")
            else:
                p = doc.add_paragraph(f"[IMAGE PLACEHOLDER: {screenshot_path}]")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add caption and explanation
            i += 1
            if i < len(report_lines) and report_lines[i].strip().startswith('Figure'):
                caption = report_lines[i].strip()
                p = doc.add_paragraph(caption)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.style = 'Caption'
                i += 1
            
            if i < len(report_lines) and report_lines[i].strip() and not report_lines[i].strip().startswith('Figure'):
                explanation = report_lines[i].strip()
                if explanation:
                    p = doc.add_paragraph(explanation)
                    p.style = 'Normal'
                i += 1
            continue
        
        # Regular text
        if line:
            if line.startswith('Section') or (line.startswith('Figure') and ':' in line):
                p = doc.add_heading(line, level=1 if line.startswith('Section') else 2)
            elif re.match(r'^\d+\.\d+', line):
                p = doc.add_heading(line, level=2)
            elif re.match(r'^\d+\.\d+\.\d+', line):
                p = doc.add_heading(line, level=3)
            else:
                p = doc.add_paragraph(line)
        
        i += 1
    
    # Add remaining screenshots that weren't in report.txt
    print(f"\nAdding remaining screenshots not in report.txt...")
    remaining = {k: v for k, v in all_screenshots.items() if k not in added_screenshots}
    
    if remaining:
        # Group by section
        doc.add_page_break()
        doc.add_heading("Additional Code Screenshots", level=1)
        
        current_section = None
        for fig_num in sorted(remaining.keys(), key=lambda x: (float(x.split('.')[0]), float(x.split('.')[1]))):
            info = remaining[fig_num]
            section = info['section']
            
            if section != current_section:
                current_section = section
                doc.add_heading(f"Section {section} - Additional Screenshots", level=2)
            
            try:
                # Add image
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run()
                run.add_picture(str(info['path']), width=Inches(6))
                screenshot_count += 1
                
                # Add caption
                caption = f"Figure {fig_num}: {info['title']}"
                p = doc.add_paragraph(caption)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.style = 'Caption'
                
                # Add explanation placeholder
                p = doc.add_paragraph(f"This screenshot shows {info['title'].lower()} implementation details.")
                p.style = 'Normal'
                
                print(f"  [{screenshot_count}] Added: {info['filename']}")
            except Exception as e:
                print(f"  Error adding {info['filename']}: {e}")
    
    doc.save(str(output_docx))
    print(f"\n✅ Word document generated: {output_docx}")
    print(f"✅ Total screenshots embedded: {screenshot_count}")
    print(f"✅ Screenshots from report.txt: {len(added_screenshots)}")
    print(f"✅ Additional screenshots added: {len(remaining)}")
    print(f"✅ You can now open this file directly in Microsoft Word!")

def main():
    report_path = Path("report.txt")
    screenshots_dir = Path("screenshots")
    output_docx = Path("report_FINAL_all_70_screenshots.docx")
    
    if not report_path.exists():
        print(f"Error: Report file not found: {report_path}")
        return
    
    print("=" * 60)
    print("Generating COMPLETE Report with ALL Screenshots")
    print("=" * 60)
    print(f"Report file: {report_path}")
    print(f"Screenshots directory: {screenshots_dir}")
    print(f"Output file: {output_docx}")
    print("-" * 60)
    
    generate_complete_report(report_path, output_docx, screenshots_dir)

if __name__ == '__main__':
    main()

