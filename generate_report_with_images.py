#!/usr/bin/env python3
"""
Generate report with embedded images for Word import
Creates both HTML (Word-compatible) and direct Word document versions
"""

import re
import base64
from pathlib import Path
from typing import Optional

def encode_image_to_base64(image_path: Path) -> Optional[str]:
    """Encode image file to base64 string"""
    if not image_path.exists():
        return None
    try:
        with open(image_path, 'rb') as f:
            image_data = f.read()
            base64_str = base64.b64encode(image_data).decode('utf-8')
            # Determine MIME type from extension
            ext = image_path.suffix.lower()
            mime_types = {
                '.png': 'image/png',
                '.jpg': 'image/jpeg',
                '.jpeg': 'image/jpeg',
                '.gif': 'image/gif'
            }
            mime_type = mime_types.get(ext, 'image/png')
            return f"data:{mime_type};base64,{base64_str}"
    except Exception as e:
        print(f"Error encoding {image_path}: {e}")
        return None

def generate_html_report(report_path: Path, output_html: Path, screenshots_dir: Path):
    """Generate HTML report with embedded images"""
    with open(report_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Find all screenshot placeholders
    pattern = r'\[INSERT SCREENSHOT: ([^\]]+)\]'
    
    def replace_screenshot(match):
        screenshot_path = match.group(1)
        full_path = screenshots_dir.parent / screenshot_path
        
        # Try to find the image
        if not full_path.exists():
            # Try alternative locations
            if 'code/' in screenshot_path:
                alt_path = screenshots_dir / Path(screenshot_path).name
                if alt_path.exists():
                    full_path = alt_path
            elif 'ui/' in screenshot_path or 'evaluation/' in screenshot_path or 'diagrams/' in screenshot_path or 'structure/' in screenshot_path:
                # Create directory if needed
                full_path.parent.mkdir(parents=True, exist_ok=True)
                # Return placeholder if file doesn't exist
                return f'<div style="border: 2px dashed #ccc; padding: 20px; text-align: center; color: #999; margin: 20px 0;"><strong>[PLACEHOLDER]</strong><br>{screenshot_path}<br><em>Image not found - please add manually</em></div>'
        
        if full_path.exists():
            base64_img = encode_image_to_base64(full_path)
            if base64_img:
                return f'<img src="{base64_img}" style="max-width: 100%; height: auto; margin: 20px 0; border: 1px solid #ddd; box-shadow: 0 2px 4px rgba(0,0,0,0.1);" alt="{Path(screenshot_path).name}" />'
        
        return f'<div style="border: 2px dashed #ccc; padding: 20px; text-align: center; color: #999; margin: 20px 0;"><strong>[PLACEHOLDER]</strong><br>{screenshot_path}<br><em>Image not found - please add manually</em></div>'
    
    # Replace screenshot placeholders with embedded images
    html_content = re.sub(pattern, replace_screenshot, content)
    
    # Convert text to HTML format
    html_content = html_content.replace('\n', '<br>\n')
    
    # Wrap in HTML document
    html_doc = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>AI-Driven Facial Recognition Attendance System - Final Report</title>
    <style>
        body {{
            font-family: 'Times New Roman', serif;
            font-size: 12pt;
            line-height: 1.5;
            max-width: 8.5in;
            margin: 0 auto;
            padding: 1in;
            background: white;
        }}
        h1 {{
            font-size: 18pt;
            font-weight: bold;
            margin-top: 24pt;
            margin-bottom: 12pt;
        }}
        h2 {{
            font-size: 14pt;
            font-weight: bold;
            margin-top: 18pt;
            margin-bottom: 9pt;
        }}
        h3 {{
            font-size: 12pt;
            font-weight: bold;
            margin-top: 12pt;
            margin-bottom: 6pt;
        }}
        p {{
            margin: 6pt 0;
            text-align: justify;
        }}
        img {{
            max-width: 100%;
            height: auto;
            margin: 12pt 0;
            display: block;
            page-break-inside: avoid;
        }}
        .figure-caption {{
            font-weight: bold;
            margin-top: 6pt;
            margin-bottom: 12pt;
            text-align: center;
            font-style: italic;
        }}
        .figure-explanation {{
            margin-bottom: 18pt;
            text-align: justify;
        }}
        pre {{
            background: #f5f5f5;
            padding: 10px;
            border: 1px solid #ddd;
            overflow-x: auto;
            font-family: 'Courier New', monospace;
            font-size: 10pt;
        }}
    </style>
</head>
<body>
{html_content}
</body>
</html>
"""
    
    with open(output_html, 'w', encoding='utf-8') as f:
        f.write(html_doc)
    
    print(f"HTML report generated: {output_html}")
    print("You can open this file in Microsoft Word and it will convert automatically!")

def generate_word_document(report_path: Path, output_docx: Path, screenshots_dir: Path):
    """Generate Word document directly using python-docx"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("python-docx not installed. Installing...")
        import subprocess
        import sys
        subprocess.check_call([sys.executable, '-m', 'pip', 'install', 'python-docx'])
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    with open(report_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # Check for screenshot placeholder
        screenshot_match = re.search(r'\[INSERT SCREENSHOT: ([^\]]+)\]', line)
        if screenshot_match:
            screenshot_path = screenshot_match.group(1)
            full_path = screenshots_dir.parent / screenshot_path
            
            # Try alternative locations
            if not full_path.exists() and 'code/' in screenshot_path:
                alt_path = screenshots_dir / Path(screenshot_path).name
                if alt_path.exists():
                    full_path = alt_path
            
            if full_path.exists():
                try:
                    # Add image
                    paragraph = doc.add_paragraph()
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph.add_run()
                    run.add_picture(str(full_path), width=Inches(6))
                    print(f"  Added image: {full_path.name}")
                except Exception as e:
                    # Add placeholder text if image can't be added
                    p = doc.add_paragraph(f"[IMAGE PLACEHOLDER: {screenshot_path}]")
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    print(f"  Could not add image {full_path}: {e}")
            else:
                # Add placeholder
                p = doc.add_paragraph(f"[IMAGE PLACEHOLDER: {screenshot_path}]")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                print(f"  Image not found: {screenshot_path}")
            
            # Add caption (next line should be the caption)
            i += 1
            if i < len(lines) and lines[i].strip().startswith('Figure'):
                caption = lines[i].strip()
                p = doc.add_paragraph(caption)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.style = 'Caption'
                i += 1
            
            # Add explanation (next line should be explanation)
            if i < len(lines) and lines[i].strip() and not lines[i].strip().startswith('Figure'):
                explanation = lines[i].strip()
                if explanation:
                    p = doc.add_paragraph(explanation)
                    p.style = 'Normal'
                i += 1
            continue
        
        # Regular text
        if line:
            # Check if it's a heading
            if line.startswith('Section') or (line.startswith('Figure') and ':' in line):
                p = doc.add_heading(line, level=1 if line.startswith('Section') else 2)
            elif re.match(r'^\d+\.\d+', line):  # Subsection like "6.1"
                p = doc.add_heading(line, level=2)
            elif re.match(r'^\d+\.\d+\.\d+', line):  # Subsubsection like "6.1.1"
                p = doc.add_heading(line, level=3)
            else:
                # Regular paragraph
                p = doc.add_paragraph(line)
        
        i += 1
    
    doc.save(str(output_docx))
    print(f"\nWord document generated: {output_docx}")
    print("You can now open this file directly in Microsoft Word!")

def main():
    import argparse
    
    parser = argparse.ArgumentParser(
        description='Generate report with embedded images for Word import'
    )
    parser.add_argument(
        '--format',
        choices=['html', 'docx', 'both'],
        default='both',
        help='Output format: html (Word-compatible), docx (direct Word), or both'
    )
    parser.add_argument(
        '--report',
        type=str,
        default='report.txt',
        help='Input report file (default: report.txt)'
    )
    parser.add_argument(
        '--screenshots-dir',
        type=str,
        default='screenshots/code',
        help='Screenshots directory (default: screenshots/code)'
    )
    
    args = parser.parse_args()
    
    report_path = Path(args.report)
    screenshots_dir = Path(args.screenshots_dir)
    
    if not report_path.exists():
        print(f"Error: Report file not found: {report_path}")
        return
    
    if args.format in ['html', 'both']:
        output_html = Path('report_with_images.html')
        generate_html_report(report_path, output_html, screenshots_dir)
    
    if args.format in ['docx', 'both']:
        output_docx = Path('report_with_images.docx')
        try:
            generate_word_document(report_path, output_docx, screenshots_dir)
        except Exception as e:
            print(f"Error generating Word document: {e}")
            print("Falling back to HTML format only")
            if args.format == 'docx':
                generate_html_report(report_path, Path('report_with_images.html'), screenshots_dir)

if __name__ == '__main__':
    main()

