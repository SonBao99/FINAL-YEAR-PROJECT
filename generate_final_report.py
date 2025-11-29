#!/usr/bin/env python3
"""
Generate final report with ALL 70 screenshots and detailed explanations
Uses explanations from SCREENSHOT_PLACEMENT_GUIDE.md
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Pre-defined detailed explanations for each screenshot
EXPLANATIONS = {
    "5.2": "This code screenshot demonstrates the Application Layer architecture (Section 5.1.2). FastAPI's asynchronous framework enables handling thousands of concurrent connections. The CORS middleware allows cross-origin requests for the web dashboard, while the ConnectionManager handles WebSocket connections for real-time updates. This design supports the scalability requirement (Section 4.3.4).",
    "5.3": "This code screenshot illustrates the database schema design (Section 5.3). Critically, the `face_encoding` field stores a 128-dimensional embedding vector as JSON (not raw images), aligning with security requirements (Section 4.3.3). The relationship to `AttendanceRecord` ensures referential integrity. This design follows Third Normal Form (3NF) normalization principles.",
    "5.4": "This code screenshot shows the temporal dimension modeling (Section 5.3). The `Session` entity includes both scheduled and actual timestamps, enabling analysis of session duration. The `is_active` flag implements a 'Soft Lock' mechanism preventing check-ins for closed sessions. The `AttendanceRecord` entity stores `confidence_score` (Euclidean distance converted to confidence), providing an audit trail for dispute resolution.",
    "5.5": "This code screenshot demonstrates the database abstraction layer (Section 5.1.3). SQLAlchemy ORM provides flexibility to operate on SQLite during development (zero-configuration) while remaining deployment-ready for PostgreSQL in production. The `get_db()` function is a FastAPI dependency that provides database sessions with automatic cleanup, ensuring proper resource management.",
    "6.1": "This screenshot shows the complete software stack used in development (Section 6.1.2). Key libraries include OpenCV for computer vision, face-recognition (dlib wrapper) for face encoding, Mediapipe for liveness detection, FastAPI for the asynchronous web framework, and SQLAlchemy for database ORM. All dependencies are version-pinned for reproducibility.",
    "6.2": "This code screenshot demonstrates the critical liveness detection implementation (Section 6.2.1). MediaPipe Face Mesh provides 468-point 3D face mesh tracking. The eye landmark indices are predefined for blink detection. Thresholds are tuned for usability: `MOVEMENT_THRESHOLD = 0.01` (sensitive to natural movement), `BLINK_THRESHOLD = 0.25` (EAR threshold), and `MIN_FRAMES_FOR_LIVE = 5` (minimum frames before determining liveness). This implements the 'Liveness First' policy mentioned in Section 6.2.1.",
    "6.3": "This code screenshot shows the mathematical foundation of blink detection (Section 6.2.1, Blink Detection). EAR measures eye openness by calculating the ratio of vertical distances (between upper and lower eyelids) to horizontal distance (eye width). When the eye closes, EAR decreases; when it opens, EAR increases. This enables detection of natural biological motion absent in static photos.",
    "6.4": "This code screenshot demonstrates the dynamic anti-spoofing mechanism (Section 6.2.1). The algorithm calculates EAR for both eyes and maintains a history buffer. A blink is detected when EAR drops below 0.25 (eye closing) and then rises above it (eye opening). The temporal analysis (comparing recent frames) distinguishes natural blinks from static photos, which cannot exhibit this temporal pattern.",
    "6.5": "This code screenshot shows the geometric analysis approach (Section 6.2.1, 3D Depth Analysis). MediaPipe provides z-coordinates (depth) for each of the 468 landmarks. Real faces have significant depth variance (nose protrudes, cheeks recede), while 2D photos on screens are geometrically flat (variance near zero). The threshold `0.0001` distinguishes real 3D faces from flat 2D representations, implementing the primary filter against presentation attacks.",
    "6.6": "This code screenshot demonstrates the multi-factor liveness detection algorithm (Section 6.2.1). The function combines movement detection, blink detection, and depth analysis into a composite liveness score. The algorithm requires at least 5 frames before making a determination, preventing false positives from transient detection failures. The status progression (CHECKING → LIVE/FAKE) provides user feedback during the verification process.",
    "6.7": "This code screenshot shows movement detection implementation (Section 6.2.1). The algorithm tracks the nose tip position across frames using a deque buffer. Movement is detected when the standard deviation of positions exceeds the threshold (0.01). This detects natural head movement that static photos cannot replicate, adding another layer of anti-spoofing protection.",
    "6.8": "This code screenshot demonstrates the spoof detection blocking mechanism (Section 6.2.2). When a face is detected as FAKE, the system displays a red bounding box and blocks recognition attempts. The status message 'Check-in blocked: FAKE face detected' provides clear feedback. This implementation ensures that spoof attempts are rejected before consuming computational resources for recognition.",
    "6.9": "This code screenshot shows liveness integration into the kiosk frame processing pipeline (Section 6.2.3). The system performs liveness detection on every frame and uses a grace period mechanism (recent liveness history) to smooth the user experience. The 'effectively_live' flag allows recognition if the face was recently LIVE, preventing flickering between states.",
    "6.10": "This code screenshot demonstrates the recognition gate mechanism (Section 6.2.3). Face recognition is only attempted when liveness is confirmed and a cooldown period has elapsed (3 seconds default). This prevents rapid-fire recognition attempts and ensures the system only processes verified live faces, optimizing computational efficiency.",
    "6.11": "This code screenshot shows the student enrollment endpoint implementation (Section 6.3.1). The endpoint decodes base64 images, converts BGR to RGB for face_recognition library, generates 128-dimensional face encodings, and stores them as JSON strings. The implementation includes error handling for invalid images and missing faces, ensuring data quality.",
    "6.12": "This code screenshot demonstrates the face recognition matching algorithm (Section 6.3.2). The system compares the detected face encoding against all enrolled students using Euclidean distance. The threshold of 0.6 determines match acceptance. The algorithm finds the best match (lowest distance) and only accepts matches below the threshold, ensuring accurate identification.",
    "6.13": "This code screenshot shows confidence score calculation (Section 6.3.2). The confidence score is calculated as `1 - distance`, converting the Euclidean distance (lower is better) to a confidence percentage (higher is better). This score is stored with each attendance record, providing an audit trail for verification accuracy and dispute resolution.",
    "6.14": "This code screenshot demonstrates the WebSocket connection manager implementation (Section 6.4.1). The ConnectionManager maintains a dictionary of active WebSocket connections organized by session ID. This allows the server to broadcast attendance updates to all connected clients for a specific session, enabling real-time dashboard updates without polling.",
    "6.15": "This code screenshot shows the WebSocket endpoint definition (Section 6.4.1). FastAPI's WebSocket support enables bidirectional communication. The endpoint accepts connections, adds them to the manager, and maintains the connection until disconnection. This implementation follows RFC 6455 WebSocket protocol standards.",
    "6.16": "This code screenshot demonstrates WebSocket broadcast functionality (Section 6.4.2). When a check-in occurs, the system broadcasts a JSON message to all connected clients for that session. The message includes student information, check-in time, and confidence score. Failed connections are automatically removed, ensuring robust real-time communication.",
    "6.17": "This code screenshot shows the frontend WebSocket client implementation (Section 6.4.3). The client establishes a WebSocket connection, handles incoming messages, and updates the UI accordingly. The implementation includes error handling and reconnection logic, ensuring reliable real-time updates even with network interruptions.",
    "6.18": "This code screenshot demonstrates kiosk-to-API communication (Section 6.5). The kiosk encodes face images to base64, sends POST requests to the check-in endpoint, and handles responses. The implementation includes timeout handling, error logging, and status message updates. This enables the edge device to communicate with the central server for face recognition processing.",
    "6.19": "This code screenshot shows the unit test suite implementation (Section 6.6). The tests use Python's unittest framework to validate AttendanceTracker functionality. Tests cover initialization, marking present/absent, roster validation, and edge cases. The test suite ensures code quality and prevents regressions during development.",
    "6.20": "This code screenshot demonstrates session management endpoints (Section 6.3.2). The endpoints allow creating sessions, listing sessions with course information, and managing session lifecycle. The implementation uses SQLAlchemy relationships to join Session and Course data, providing complete session information in a single API call.",
    "6.21": "This code screenshot shows the manual check-in endpoint implementation (Section 6.3.3). This endpoint allows administrative override for edge cases where face recognition fails. The implementation validates session and student existence, prevents duplicate check-ins, and broadcasts WebSocket updates, maintaining consistency with automated check-ins.",
    "6.22": "This code screenshot demonstrates kiosk initialization and camera setup (Section 6.2.4). The kiosk initializes OpenCV VideoCapture, configures camera properties (640x480, 30 FPS), and loads the Haar Cascade classifier for face detection. This setup ensures optimal camera performance for face detection and recognition.",
    "6.23": "This code screenshot shows Chart.js integration for attendance visualization (Section 6.7.1). The implementation creates doughnut charts showing present/absent distribution. Chart.js provides responsive, interactive visualizations that update in real-time as attendance data changes, enhancing the dashboard's visual appeal and usability.",
    "6.24": "This code screenshot demonstrates the polling fallback mechanism (Section 6.4.4). If WebSocket connection fails, the system falls back to HTTP polling every 30 seconds. This ensures the dashboard remains functional even with WebSocket connectivity issues, providing a robust user experience.",
    "6.25": "This code screenshot shows comprehensive error handling in face recognition (Section 6.3.4). The implementation handles invalid images, missing faces, session validation errors, and network failures. Each error type provides specific error messages, enabling debugging and user feedback. This robust error handling ensures system reliability.",
    "6.26": "This code screenshot demonstrates database session dependency injection (Section 6.3.5). FastAPI's dependency injection system provides database sessions to endpoints. The `get_db()` function yields a session and ensures cleanup in a finally block, preventing connection leaks. This pattern is used throughout all database operations.",
    "6.27": "This code screenshot shows WebSocket reconnection logic (Section 6.4.5). The frontend implements exponential backoff reconnection with a maximum of 5 attempts. The delay increases exponentially (1s, 2s, 4s, 8s, 16s) up to 30 seconds maximum. This prevents overwhelming the server while ensuring eventual reconnection.",
    "6.28": "This code screenshot demonstrates face encoding storage format (Section 6.3.6). Face encodings (128-dimensional NumPy arrays) are converted to JSON strings before database storage. This format allows efficient storage and retrieval while maintaining compatibility with JSON-based APIs. The encoding can be deserialized back to NumPy arrays for distance calculations.",
    "6.35": "This code screenshot shows the health check endpoint and startup event handler. The health check endpoint allows monitoring tools to verify API availability, while the startup event ensures database tables are created before handling requests. This demonstrates proper application lifecycle management in FastAPI.",
    "6.36": "These endpoints control session lifecycle, setting the `is_active` flag and recording actual start/end times. The implementation demonstrates proper state management and timestamp tracking for audit purposes. The `actual_start` and `actual_end` fields allow comparison with scheduled times for attendance analysis.",
    "6.37": "This endpoint retrieves all attendance records for a session, using SQLAlchemy relationships to join Student data. The response model demonstrates nested Pydantic models (`AttendanceResponse` containing `StudentResponse`), ensuring type safety and automatic API documentation. This pattern allows efficient data retrieval while maintaining clean separation of concerns.",
    "6.38": "This script demonstrates batch enrollment capabilities, processing multiple images per student. The function reads images using OpenCV, encodes them to base64 for API transmission, and handles errors gracefully. This utility script enables efficient bulk enrollment operations, supporting the administrative workflow requirement.",
    "6.39": "This main function demonstrates robust input handling, supporting directory-based, comma-separated, or interactive image path entry. The script validates paths, enforces the 3-5 image requirement, and provides progress feedback. This implementation supports the multiple reference images requirement and demonstrates error handling and user feedback best practices.",
    "6.40": "This function demonstrates the course creation workflow, handling optional description fields and providing clear success/error feedback. The function returns the course ID for use in session creation, demonstrating proper workflow chaining. This supports the course management requirement.",
    "6.41": "This function creates attendance sessions with scheduled times and optional room location. The datetime objects are serialized to ISO format for JSON transmission, demonstrating proper API communication patterns. The function returns the session ID for subsequent operations like starting the session.",
    "6.42": "These methods form the core of the attendance tracking system, using Python sets for efficient membership testing. The `mark_present` and `mark_absent` methods validate against the roster, preventing invalid operations. The `get_absent_students` method uses set difference for efficient computation, demonstrating algorithmic efficiency.",
    "6.43": "These methods provide flexible data persistence, supporting both JSON (for programmatic use) and CSV (for spreadsheet compatibility) formats. The `from_dict` method includes validation to ensure loaded present students are in the roster, maintaining data integrity. The CSV format uses binary encoding ('1'/'0') for present/absent status, enabling easy analysis in spreadsheet applications.",
    "6.44": "This method allows the kiosk operator to select an active session interactively. The function fetches all sessions from the API, filters for active ones, and presents them in a numbered list. This demonstrates proper error handling for network failures and invalid user input, ensuring robust operation in production environments.",
    "6.45": "This method saves face snapshots for audit purposes, using UTC timestamps and sanitized labels in filenames. The optional feature can be toggled via `save_snapshots` flag, allowing operators to enable/disable based on privacy requirements. The snapshots provide evidence for attendance disputes and system debugging.",
    "6.46": "This function demonstrates the frontend's data fetching and rendering logic. It handles empty states, formats timestamps for display, and updates multiple UI components (list, statistics, charts) after fetching data. The function uses async/await for clean asynchronous code and includes error handling with user-friendly error messages.",
    "6.47": "This function demonstrates data export capabilities, converting attendance records to Excel format using the SheetJS (XLSX) library. The function formats data appropriately, creates a workbook with a named sheet, and generates a filename with session name and date. This supports the reporting requirement and enables offline analysis.",
    "6.48": "This code implements dark mode functionality using CSS custom properties (CSS variables) and the `data-theme` attribute. The theme preference is persisted in localStorage, ensuring the user's choice persists across sessions. The implementation demonstrates modern web development practices and improves user experience through accessibility features.",
    "6.49": "This function creates user-friendly toast notifications with different types (success, error, warning, info) and corresponding icons. The notifications auto-dismiss after a configurable duration with smooth animations. This provides immediate feedback for user actions, improving the overall user experience and supporting the usability requirement.",
    "6.50": "This function renders session cards with dynamic status determination based on `is_active` flag and scheduled times. The function includes conditional rendering for optional fields (room location) and disables action buttons based on session state. This demonstrates proper state management and user interface design, ensuring users can only perform valid actions.",
    "6.51": "This file lists all project dependencies with minimum version requirements. The dependencies cover computer vision (OpenCV, face_recognition, MediaPipe), web framework (FastAPI, uvicorn), database (SQLAlchemy, psycopg2, pymongo), and utility libraries. Version constraints ensure reproducible builds and compatibility across development environments.",
    "6.52": "This endpoint demonstrates SQLAlchemy relationship navigation, joining Session and Course data through the foreign key relationship. The function handles nullable relationships and serializes datetime objects to ISO format strings for JSON compatibility. This pattern provides complete session information in a single API call, reducing frontend complexity.",
    "6.53": "This main loop demonstrates robust camera handling with failure recovery, frame processing with error handling, and interactive keyboard controls. The loop supports session selection ('s'), session refresh ('r'), snapshot toggle ('t'), and graceful shutdown ('q'). The implementation includes proper resource cleanup (camera release, window destruction) ensuring no resource leaks.",
    "6.54": "This function computes real-time attendance statistics from records, calculating total students, present/absent counts, and attendance rate percentage. The function uses JavaScript array filtering and updates DOM elements dynamically, demonstrating reactive UI updates based on data changes.",
    "6.55": "This async function handles session creation from the frontend, collecting form data, converting dates to ISO format, and sending a POST request to the API. The function includes error handling, success notifications, and UI updates, demonstrating proper async/await patterns and user feedback mechanisms.",
    "6.56": "This function handles student enrollment with multiple photos, validating minimum photo count, converting files to base64, and sending sequential API requests. The implementation demonstrates batch processing, file handling, and proper error handling for the first enrollment attempt.",
    "6.57": "These functions handle file preview and conversion: `previewPhotos` creates thumbnail previews using FileReader API, `removePhoto` allows removing individual photos, and `fileToBase64` converts files to base64 strings for API transmission. The implementation demonstrates modern browser file handling APIs and Promise-based async operations.",
    "6.58": "These functions control session lifecycle, calling API endpoints to start/stop sessions and updating the UI accordingly. The implementation demonstrates RESTful API patterns, error handling, and coordinated UI updates across multiple components (sessions list, controls, notifications).",
    "6.59": "These functions implement client-side filtering with multiple criteria: text search, status filtering, and course filtering. The `filterSessions` function demonstrates complex boolean logic for status determination based on dates and active flags, while `filterAttendance` uses DOM manipulation for real-time filtering without re-rendering.",
    "6.60": "This function implements client-side sorting by name, time, or status. The function uses dataset attributes to maintain references to original data, sorts DOM elements, and re-appends them in sorted order. This demonstrates efficient DOM manipulation without full re-rendering.",
    "6.61": "This function generates CSV data from attendance records, creates a Blob object, and triggers browser download using URL.createObjectURL. The implementation demonstrates proper CSV formatting with quoted fields, date formatting, and memory cleanup via URL.revokeObjectURL.",
    "6.62": "This function uses the SheetJS (XLSX) library to create Excel workbooks, converting attendance data to worksheet format. The implementation demonstrates library integration, workbook creation, and file generation with proper naming conventions.",
    "6.63": "These functions enable manual attendance entry: opening the modal, loading student options, and submitting entries. The implementation demonstrates form validation, date handling, student lookup, and API integration for administrative override capabilities.",
    "6.64": "This function generates comprehensive attendance reports by fetching sessions, filtering by date range, aggregating attendance data across multiple sessions, and calculating statistics. The implementation demonstrates batch API calls, data aggregation, and chart updates.",
    "6.65": "This function computes True Positives, False Positives, False Negatives, and True Negatives by comparing ground truth and predicted values. The implementation handles edge cases (empty values) and builds a confusion matrix dataset, demonstrating proper evaluation metric calculation for face recognition systems.",
    "6.66": "This function calculates standard biometric evaluation metrics: Precision, Recall, F1-score, Accuracy, False Acceptance Rate (FAR), and False Rejection Rate (FRR). The implementation includes division-by-zero protection and demonstrates proper metric calculation for face recognition evaluation.",
    "6.67": "This function computes advanced biometric metrics: ROC curve, AUC (Area Under Curve), and EER (Equal Error Rate). The implementation uses scikit-learn for ROC computation and calculates EER by finding the threshold where False Positive Rate equals False Negative Rate, demonstrating sophisticated evaluation techniques.",
    "6.68": "These test methods demonstrate unit testing patterns: `setUp` provides test fixtures, `test_initialization` verifies initial state, and `test_mark_present` tests core functionality. The tests use unittest assertions (`assertEqual`, `assertIn`, `assertNotIn`) to validate expected behavior.",
    "6.69": "These models define Course and Lecturer entities with proper SQLAlchemy column definitions, constraints (unique, nullable), and relationships. The Course model includes a one-to-many relationship with Sessions, while the Lecturer model includes authentication fields (hashed_password) for future authentication implementation.",
    "6.70": "These endpoints provide course management functionality: GET retrieves all active courses with filtering, and POST creates new courses using Pydantic models for validation. The implementation demonstrates standard CRUD patterns, database session management, and proper HTTP status codes.",
    "6.71": "This function creates and updates multiple Chart.js visualizations: a line chart for attendance trends over time and a pie chart for status distribution. The implementation demonstrates data aggregation, chart lifecycle management (destroying old charts before creating new ones), and Chart.js configuration.",
    "6.72": "This function implements web-based face recognition check-in by capturing video frames to canvas, converting to base64, and sending to the API. The function handles confidence thresholds, success/error states, and provides user feedback with colored status messages and auto-dismiss timers.",
}

def find_all_screenshots(screenshots_dir: Path):
    """Find all screenshot files and organize by figure number"""
    screenshots = {}
    code_dir = screenshots_dir / 'code'
    
    if code_dir.exists():
        for png_file in sorted(code_dir.glob('fig-*.png')):
            match = re.search(r'fig-(\d+\.\d+)-(.+)\.png', png_file.name)
            if match:
                fig_num = match.group(1)
                title = match.group(2).replace('-', ' ')
                screenshots[fig_num] = {
                    'path': png_file,
                    'filename': png_file.name,
                    'title': title,
                    'section': fig_num.split('.')[0]
                }
    
    return screenshots

def get_explanation(fig_num: str, title: str) -> str:
    """Get explanation for a screenshot"""
    if fig_num in EXPLANATIONS:
        return EXPLANATIONS[fig_num]
    
    # Generate default explanation
    section = fig_num.split('.')[0]
    return f"This code screenshot demonstrates {title.lower()} implementation (Section {section}). The code shows key functionality and design patterns used in the system. This implementation supports the system's requirements and architectural goals."

def generate_complete_report(report_path: Path, output_docx: Path, screenshots_dir: Path):
    """Generate Word document with ALL screenshots and detailed explanations"""
    
    # Find all screenshots
    all_screenshots = find_all_screenshots(screenshots_dir)
    print(f"Found {len(all_screenshots)} code screenshots")
    
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Read report content
    with open(report_path, 'r', encoding='utf-8') as f:
        report_lines = f.readlines()
    
    # Track which screenshots we've added
    added_screenshots = set()
    screenshot_count = 0
    
    # Process report and add screenshots where referenced
    i = 0
    while i < len(report_lines):
        line = report_lines[i].strip()
        
        # Check for screenshot placeholder
        screenshot_match = re.search(r'\[INSERT SCREENSHOT: ([^\]]+)\]', line)
        if screenshot_match:
            screenshot_path = screenshot_match.group(1)
            
            # Try to find the screenshot
            screenshot_file = None
            fig_num = None
            
            # Extract figure number from path
            match = re.search(r'fig-(\d+\.\d+)-', screenshot_path)
            if match:
                fig_num = match.group(1)
                if fig_num in all_screenshots:
                    screenshot_file = all_screenshots[fig_num]['path']
                    added_screenshots.add(fig_num)
            
            # Try alternative: just filename
            if not screenshot_file:
                filename = Path(screenshot_path).name
                for fnum, info in all_screenshots.items():
                    if info['filename'] == filename:
                        screenshot_file = info['path']
                        fig_num = fnum
                        added_screenshots.add(fnum)
                        break
            
            # Try alternative: direct path
            if not screenshot_file:
                full_path = screenshots_dir.parent / screenshot_path
                if full_path.exists():
                    screenshot_file = full_path
                    match = re.search(r'fig-(\d+\.\d+)-', full_path.name)
                    if match:
                        fig_num = match.group(1)
            
            if screenshot_file and screenshot_file.exists():
                try:
                    # Add image
                    paragraph = doc.add_paragraph()
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph.add_run()
                    run.add_picture(str(screenshot_file), width=Inches(6))
                    screenshot_count += 1
                    print(f"  [{screenshot_count}] Added: {screenshot_file.name}")
                    
                    # Add caption
                    if fig_num and fig_num in all_screenshots:
                        title = all_screenshots[fig_num]['title']
                        caption = f"Figure {fig_num}: {title}"
                    else:
                        caption = f"Figure {screenshot_file.stem}"
                    
                    p = doc.add_paragraph(caption)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.style = 'Caption'
                    
                    # Add explanation from report or predefined
                    explanation = None
                    if i + 1 < len(report_lines):
                        next_line = report_lines[i + 1].strip()
                        if next_line and not next_line.startswith('Figure'):
                            explanation = next_line
                            i += 1
                    
                    if not explanation and fig_num:
                        explanation = get_explanation(fig_num, 
                                                    all_screenshots.get(fig_num, {}).get('title', ''))
                    
                    if explanation:
                        p = doc.add_paragraph(explanation)
                        p.style = 'Normal'
                        p.paragraph_format.space_after = Pt(12)
                    
                except Exception as e:
                    p = doc.add_paragraph(f"[IMAGE ERROR: {e}]")
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    print(f"  Error adding {screenshot_file.name}: {e}")
            else:
                p = doc.add_paragraph(f"[IMAGE PLACEHOLDER: {screenshot_path}]")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Skip caption and explanation if already processed
            i += 1
            if i < len(report_lines) and report_lines[i].strip().startswith('Figure'):
                i += 1
            if i < len(report_lines) and report_lines[i].strip() and not report_lines[i].strip().startswith('Figure'):
                i += 1
            continue
        
        # Regular text
        if line:
            if line.startswith('Section') or (line.startswith('Figure') and ':' in line):
                p = doc.add_heading(line, level=1 if line.startswith('Section') else 2)
            elif re.match(r'^\d+\.\d+', line):
                p = doc.add_heading(line, level=2)
            elif re.match(r'^\d+\.\d+\.\d+', line):
                p = doc.add_heading(line, level=3)
            else:
                p = doc.add_paragraph(line)
        
        i += 1
    
    # Add remaining screenshots that weren't in report.txt
    print(f"\nAdding remaining screenshots with detailed explanations...")
    remaining = {k: v for k, v in all_screenshots.items() if k not in added_screenshots}
    
    if remaining:
        # Group by section
        doc.add_page_break()
        doc.add_heading("Additional Code Screenshots", level=1)
        doc.add_paragraph("The following screenshots provide additional implementation details not explicitly referenced in the main report text. Each screenshot includes detailed explanations of the code's functionality, design patterns, and contribution to the overall system architecture.")
        doc.add_paragraph("")
        
        current_section = None
        for fig_num in sorted(remaining.keys(), key=lambda x: (float(x.split('.')[0]), float(x.split('.')[1]))):
            info = remaining[fig_num]
            section = info['section']
            
            if section != current_section:
                current_section = section
                doc.add_heading(f"Section {section} - Additional Screenshots", level=2)
            
            try:
                # Add image
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run()
                run.add_picture(str(info['path']), width=Inches(6))
                screenshot_count += 1
                
                # Add caption
                caption = f"Figure {fig_num}: {info['title']}"
                p = doc.add_paragraph(caption)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.style = 'Caption'
                
                # Add detailed explanation
                explanation = get_explanation(fig_num, info['title'])
                p = doc.add_paragraph(explanation)
                p.style = 'Normal'
                p.paragraph_format.space_after = Pt(12)
                
                print(f"  [{screenshot_count}] Added with explanation: {info['filename']}")
            except Exception as e:
                print(f"  Error adding {info['filename']}: {e}")
    
    doc.save(str(output_docx))
    print(f"\n✅ Word document generated: {output_docx}")
    print(f"✅ Total screenshots embedded: {screenshot_count}")
    print(f"✅ Screenshots from report.txt: {len(added_screenshots)}")
    print(f"✅ Additional screenshots added: {len(remaining)}")
    print(f"✅ All screenshots include detailed explanations!")
    print(f"✅ You can now open this file directly in Microsoft Word!")

def main():
    report_path = Path("report.txt")
    screenshots_dir = Path("screenshots")
    output_docx = Path("report_FINAL_with_explanations.docx")
    
    if not report_path.exists():
        print(f"Error: Report file not found: {report_path}")
        return
    
    print("=" * 70)
    print("Generating Report with ALL Screenshots and Detailed Explanations")
    print("=" * 70)
    print(f"Report file: {report_path}")
    print(f"Screenshots directory: {screenshots_dir}")
    print(f"Output file: {output_docx}")
    print("-" * 70)
    
    generate_complete_report(report_path, output_docx, screenshots_dir)

if __name__ == '__main__':
    main()

